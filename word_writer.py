from typing import Dict, List, Any

def write_summary_to_word(summary_data: Dict[str, Any], output_path: str = 'summary_output.docx') -> None:
    """
    Write the summary data to a Word document in full sentences using python-docx.
    Args:
        summary_data: The summary dictionary parsed from JSON.
        output_path: The path to save the Word document.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Please install python-docx: pip install python-docx")

    doc = Document()
    doc.add_heading('Call Summary', 0)

    # Extract call info, highlights, and action items from the summary data
    call_info = summary_data.get('call_info', {})
    highlights = summary_data.get('conversation_highlights', [])
    action_items = summary_data.get('action_items', [])

    # Add call info as a summary sentence
    customer = call_info.get('customer_name', 'Unknown Customer')
    agent = call_info.get('agent_name', 'Unknown Agent')
    date = call_info.get('date', None)
    duration = call_info.get('duration', None)
    info_sentence = f"Customer {customer} spoke with agent {agent}."
    if date:
        info_sentence += f" The call took place on {date}."
    if duration:
        info_sentence += f" The call lasted approximately {duration}."
    doc.add_paragraph(info_sentence)

    # Add conversation highlights as bullet points
    if highlights:
        doc.add_heading('Conversation Highlights', level=1)
        for h in highlights:
            doc.add_paragraph(h, style='List Bullet')

    # Add action items as numbered sentences, one per line
    if action_items:
        doc.add_heading('Action Items', level=1)
        for idx, item in enumerate(action_items, 1):
            task = item.get('task', 'No task specified')
            assigned = item.get('assigned_to', 'Unassigned')
            due = item.get('due', None)
            if due:
                sentence = f"{idx}. {assigned} needs to {task} by {due}."
            else:
                sentence = f"{idx}. {assigned} needs to {task}."
            doc.add_paragraph(sentence)

    doc.save(output_path)
    print(f"Summary data written to {output_path} (Word document)")
